import html
import inspect
import json
import os
import re
import tempfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import quote

import gradio as gr
import matplotlib
import requests
from bs4 import BeautifulSoup
from docx import Document
from huggingface_hub import InferenceClient
from pypdf import PdfReader


matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402

APP_NAME = "OpenStudy"
QWEN_MODEL_ID = "Qwen/Qwen3.6-27B"
MAX_ANALYSIS_CHARS = 180_000
MAX_QWEN_CONTEXT_CHARS = 18_000
REQUEST_TIMEOUT = 12
HEADERS = {
    "User-Agent": "OpenStudy/0.1 (open-source research-bias-screening; https://huggingface.co/spaces)"
}
# --- OpenStudy design tokens ---------------------------------------------
BRAND = "#0e6b5f"
BRAND_DEEP = "#0a554b"
BRAND_BRIGHT = "#2fa08f"
PAPER = "#f6f4ee"
CARD = "#fffdf8"
INK = "#20312d"
INK_SOFT = "#5b6b66"
BORDER = "#e4dfd3"
WARN = "#b45309"
RISK = "#b91c1c"

THEME = gr.themes.Soft(
    primary_hue="teal",
    secondary_hue="emerald",
    neutral_hue="stone",
    font=[gr.themes.GoogleFont("Inter"), "ui-sans-serif", "system-ui", "sans-serif"],
    font_mono=[gr.themes.GoogleFont("IBM Plex Mono"), "ui-monospace", "SFMono-Regular", "monospace"],
).set(
    body_background_fill=PAPER,
    body_background_fill_dark="#101614",
    body_text_color=INK,
    body_text_color_dark="#e9e7df",
    body_text_color_subdued=INK_SOFT,
    body_text_color_subdued_dark="#9fafa9",
    background_fill_primary=CARD,
    background_fill_primary_dark="#19221f",
    background_fill_secondary=PAPER,
    background_fill_secondary_dark="#141c19",
    border_color_primary=BORDER,
    border_color_primary_dark="#2b3733",
    block_background_fill=CARD,
    block_background_fill_dark="#19221f",
    block_border_color=BORDER,
    block_border_color_dark="#2b3733",
    block_shadow="0 1px 2px rgba(32, 49, 45, 0.05)",
    block_title_text_color=INK,
    block_title_text_color_dark="#e9e7df",
    block_label_text_color=INK_SOFT,
    block_label_text_color_dark="#9fafa9",
    block_label_background_fill=PAPER,
    block_label_background_fill_dark="#141c19",
    input_background_fill="#ffffff",
    input_background_fill_dark="#101614",
    input_border_color="#d8d2c4",
    input_border_color_dark="#33403b",
    button_primary_background_fill=BRAND,
    button_primary_background_fill_hover=BRAND_DEEP,
    button_primary_background_fill_dark=BRAND,
    button_primary_background_fill_hover_dark="#128172",
    button_primary_text_color="#ffffff",
    button_primary_text_color_dark="#ffffff",
    button_secondary_background_fill="#edebe0",
    button_secondary_background_fill_hover="#e3e0d2",
    button_secondary_background_fill_dark="#243029",
    button_secondary_background_fill_hover_dark="#2c3a32",
    button_secondary_text_color=INK,
    button_secondary_text_color_dark="#e9e7df",
    table_even_background_fill="#fbfaf4",
    table_even_background_fill_dark="#1b2421",
    table_odd_background_fill=CARD,
    table_odd_background_fill_dark="#19221f",
    link_text_color=BRAND,
    link_text_color_dark="#3db2a0",
    color_accent_soft="#e3efec",
    color_accent_soft_dark="#16302b",
    slider_color=BRAND,
    slider_color_dark=BRAND_BRIGHT,
)
STUDY_TYPE_CHOICES = [
    "Auto-detect",
    "Clinical trial",
    "Observational study",
    "Systematic review or meta-analysis",
    "Qualitative study",
    "Survey study",
    "Animal or preclinical study",
    "Model or algorithm study",
]


@dataclass(frozen=True)
class Criterion:
    category: str
    label: str
    positive: tuple[str, ...]
    recommendation: str
    concern: tuple[str, ...] = ()
    applies_to: tuple[str, ...] = ()
    excludes: tuple[str, ...] = ()


@dataclass(frozen=True)
class Benchmark:
    name: str
    description: str
    criteria: tuple[str, ...]
    applies_to: tuple[str, ...] = ()
    excludes: tuple[str, ...] = ()


CRITERIA: tuple[Criterion, ...] = (
    Criterion(
        "Design and protocol",
        "Study design is named",
        (
            r"\brandomi[sz]ed\b",
            r"\bclinical trial\b",
            r"\bcohort\b",
            r"\bcase[- ]control\b",
            r"\bcross[- ]sectional\b",
            r"\bsurvey\b",
            r"\bqualitative\b",
            r"\bsystematic review\b",
            r"\bmeta[- ]analysis\b",
            r"\bprospective\b",
            r"\bretrospective\b",
        ),
        "State the study design clearly enough for readers to judge what claims are supported.",
    ),
    Criterion(
        "Design and protocol",
        "Protocol or preregistration is reported",
        (
            r"\bpre[- ]registered\b",
            r"\bpreregistration\b",
            r"\bregistered\b",
            r"\bclinicaltrials\.gov\b",
            r"\bprotocol\b",
            r"\bosf\.io\b",
            r"\bprospero\b",
        ),
        "Report a protocol, registry entry, or explain why preregistration was not possible.",
    ),
    Criterion(
        "Design and protocol",
        "Primary outcomes or endpoints are specified",
        (
            r"\bprimary outcome\b",
            r"\bprimary endpoint\b",
            r"\bsecondary outcome\b",
            r"\bend ?point\b",
            r"\boutcome measures?\b",
        ),
        "Identify primary and secondary outcomes before presenting results.",
    ),
    Criterion(
        "Sampling and participants",
        "Eligibility criteria are described",
        (
            r"\binclusion criteria\b",
            r"\bexclusion criteria\b",
            r"\beligib(le|ility)\b",
            r"\bwere included\b",
            r"\bwere excluded\b",
        ),
        "Describe who could enter the study and who was excluded.",
    ),
    Criterion(
        "Sampling and participants",
        "Recruitment source or study period is described",
        (
            r"\brecruited\b",
            r"\benrolled\b",
            r"\bconsecutive\b",
            r"\bstudy period\b",
            r"\bbetween [a-z]+ \d{4}\b",
            r"\bfrom \d{4} to \d{4}\b",
            r"\bdata were collected\b",
        ),
        "Report where, when, and how participants or records were selected.",
    ),
    Criterion(
        "Sampling and participants",
        "Sample size or power rationale is reported",
        (
            r"\bsample size\b",
            r"\bpower (analysis|calculation)\b",
            r"\bstatistical power\b",
            r"\bminimum detectable\b",
        ),
        "Provide a sample-size rationale, power analysis, or precision target.",
    ),
    Criterion(
        "Sampling and participants",
        "Participant characteristics are reported",
        (
            r"\bbaseline characteristics\b",
            r"\bdemographic",
            r"\bage\b",
            r"\bsex\b",
            r"\bgender\b",
            r"\brace\b",
            r"\bethnicity\b",
        ),
        "Report participant characteristics relevant to fairness and generalizability.",
    ),
    Criterion(
        "Bias controls",
        "Random allocation is reported",
        (
            r"\brandomi[sz]ed\b",
            r"\brandom allocation\b",
            r"\ballocation sequence\b",
            r"\bblock random",
            r"\bstratified random",
        ),
        "Describe the allocation method and who generated the sequence.",
        applies_to=("clinical_trial", "animal"),
    ),
    Criterion(
        "Bias controls",
        "Blinding or masking is reported",
        (
            r"\bdouble[- ]blind\b",
            r"\bsingle[- ]blind\b",
            r"\btriple[- ]blind\b",
            r"\bblind(ed|ing)\b",
            r"\bmask(ed|ing)\b",
        ),
        "State whether participants, investigators, outcome assessors, or analysts were blinded.",
        concern=(r"\bopen[- ]label\b", r"\bunblinded\b", r"\bnot blinded\b"),
        applies_to=("clinical_trial", "animal"),
    ),
    Criterion(
        "Bias controls",
        "Comparator or control condition is described",
        (
            r"\bcontrol group\b",
            r"\bcontrol condition\b",
            r"\bcomparator\b",
            r"\bplacebo\b",
            r"\busual care\b",
            r"\bmatched controls?\b",
        ),
        "Describe the comparator, control, or counterfactual condition.",
        excludes=("review", "qualitative"),
    ),
    Criterion(
        "Bias controls",
        "Confounding is addressed",
        (
            r"\bconfound",
            r"\badjusted for\b",
            r"\bmultivariable\b",
            r"\bmultivariate\b",
            r"\bpropensity score\b",
            r"\binverse probability\b",
            r"\bstratified analysis\b",
        ),
        "Explain likely confounders and how the analysis reduces confounding.",
        applies_to=("observational", "survey", "unknown"),
    ),
    Criterion(
        "Bias controls",
        "Missing data, attrition, or follow-up is addressed",
        (
            r"\bmissing data\b",
            r"\blost to follow[- ]up\b",
            r"\battrition\b",
            r"\bwithdrawn\b",
            r"\bimputation\b",
            r"\bcomplete case\b",
        ),
        "Report missing data, exclusions after enrollment, attrition, and handling methods.",
    ),
    Criterion(
        "Bias controls",
        "Measurement validity or reliability is addressed",
        (
            r"\bvalidated\b",
            r"\bvalidation\b",
            r"\breliability\b",
            r"\binter[- ]rater\b",
            r"\bintra[- ]class\b",
            r"\bcalibrat",
            r"\bstandardi[sz]ed\b",
        ),
        "Describe validated instruments, calibration, or reliability checks for key measures.",
    ),
    Criterion(
        "Analysis",
        "Statistical methods are named",
        (
            r"\bregression\b",
            r"\banova\b",
            r"\bt[- ]test\b",
            r"\bchi[- ]square\b",
            r"\bhazard ratio\b",
            r"\bodds ratio\b",
            r"\brisk ratio\b",
            r"\bbayesian\b",
            r"\bmodel\b",
        ),
        "Name the statistical methods and connect them to the hypotheses or outcomes.",
    ),
    Criterion(
        "Analysis",
        "Uncertainty is reported",
        (
            r"\bconfidence interval\b",
            r"\bcredible interval\b",
            r"\bstandard error\b",
            r"\bstandard deviation\b",
            r"\beffect size\b",
            r"\bp ?[<=>]\s?0\.",
            r"\bp-value\b",
        ),
        "Report effect sizes and uncertainty, not only direction or significance.",
    ),
    Criterion(
        "Analysis",
        "Sensitivity, subgroup, or robustness checks are reported",
        (
            r"\bsensitivity analysis\b",
            r"\brobustness\b",
            r"\bsubgroup analysis\b",
            r"\bmultiple comparisons\b",
            r"\bfalse discovery\b",
            r"\bbonferroni\b",
        ),
        "Show whether conclusions hold under reasonable alternative analyses.",
    ),
    Criterion(
        "Transparency",
        "Data or code availability is reported",
        (
            r"\bdata availability\b",
            r"\bavailable upon request\b",
            r"\bopen data\b",
            r"\bcode (is )?available\b",
            r"\bgithub\.com\b",
            r"\brepository\b",
            r"\bsupplementary data\b",
        ),
        "Provide data, code, materials, or a justified access limitation.",
    ),
    Criterion(
        "Ethics",
        "Ethics approval or oversight is reported",
        (
            r"\binstitutional review board\b",
            r"\birb\b",
            r"\bethics committee\b",
            r"\bethical approval\b",
            r"\bethics approval\b",
            r"\bapproved by\b",
            r"\bdeclaration of helsinki\b",
        ),
        "Report the oversight body, approval number when available, or exemption rationale.",
        excludes=("review",),
    ),
    Criterion(
        "Ethics",
        "Consent process is reported",
        (
            r"\binformed consent\b",
            r"\bconsent (was )?obtained\b",
            r"\bwritten consent\b",
            r"\bwaiver of consent\b",
            r"\bassent\b",
        ),
        "State how consent was obtained or why consent was waived.",
        excludes=("review", "animal"),
    ),
    Criterion(
        "Ethics",
        "Privacy, confidentiality, or safety safeguards are reported",
        (
            r"\bconfidential",
            r"\banonymi[sz]ed\b",
            r"\bde[- ]identified\b",
            r"\bprivacy\b",
            r"\badverse events?\b",
            r"\bsafety monitoring\b",
            r"\bdata protection\b",
        ),
        "Describe participant privacy protections and safety monitoring where relevant.",
        excludes=("review", "animal"),
    ),
    Criterion(
        "Ethics",
        "Animal welfare oversight is reported",
        (
            r"\biacuc\b",
            r"\banimal care\b",
            r"\banimal ethics\b",
            r"\banimal welfare\b",
            r"\barrive\b",
        ),
        "Report animal-care approval and welfare safeguards.",
        applies_to=("animal",),
    ),
    Criterion(
        "Funding and conflicts",
        "Funding source is reported",
        (
            r"\bfunding\b",
            r"\bfunded by\b",
            r"\bgrant\b",
            r"\bsponsor(ed|ship)?\b",
            r"\bfinancial support\b",
        ),
        "Report funding sources or state that no external funding was received.",
    ),
    Criterion(
        "Funding and conflicts",
        "Conflicts or competing interests are reported",
        (
            r"\bconflicts? of interest\b",
            r"\bcompeting interests?\b",
            r"\bfinancial interests?\b",
            r"\bno competing interests\b",
            r"\bno conflicts?\b",
        ),
        "Disclose conflicts of interest or explicitly state there were none.",
    ),
    Criterion(
        "Funding and conflicts",
        "Funder or sponsor role is clarified",
        (
            r"\bfunder(s)? had no role\b",
            r"\bsponsor(s)? had no role\b",
            r"\bindependent of the funder\b",
            r"\brole of the funding source\b",
        ),
        "Clarify whether funders influenced design, analysis, interpretation, or publication.",
    ),
    Criterion(
        "Interpretation",
        "Limitations are discussed",
        (
            r"\blimitations?\b",
            r"\bstrengths and limitations\b",
            r"\bexternal validity\b",
            r"\bgeneralizability\b",
            r"\bmay not generalize\b",
        ),
        "Discuss limitations, generalizability, and remaining uncertainty.",
    ),
    Criterion(
        "Interpretation",
        "Causal claims are appropriately constrained",
        (
            r"\bassociated with\b",
            r"\bassociation between\b",
            r"\bcorrelat",
            r"\bcannot infer caus",
            r"\bnot establish caus",
            r"\bobservational\b",
        ),
        "Avoid causal language unless the design supports causal inference.",
        concern=(r"\bcauses?\b", r"\bcaused\b", r"\bleads to\b", r"\bprevents?\b"),
        applies_to=("observational", "survey", "unknown"),
    ),
    Criterion(
        "Review methods",
        "Search strategy and databases are reported",
        (
            r"\bsearch strategy\b",
            r"\bpubmed\b",
            r"\bmedline\b",
            r"\bembase\b",
            r"\bweb of science\b",
            r"\bscopus\b",
            r"\bcochrane\b",
        ),
        "Report databases, search dates, and enough query detail to reproduce the review.",
        applies_to=("review",),
    ),
    Criterion(
        "Review methods",
        "Risk-of-bias or quality appraisal is reported",
        (
            r"\brisk of bias\b",
            r"\bquality assessment\b",
            r"\bcochrane risk\b",
            r"\brobins\b",
            r"\bnewcastle[- ]ottawa\b",
            r"\bgrade\b",
        ),
        "Assess included-study quality or risk of bias with an appropriate tool.",
        applies_to=("review",),
    ),
    Criterion(
        "Review methods",
        "Publication bias is considered",
        (
            r"\bpublication bias\b",
            r"\bfunnel plot\b",
            r"\begger",
            r"\btrim and fill\b",
        ),
        "Assess publication bias or explain why it could not be assessed.",
        applies_to=("review",),
    ),
    Criterion(
        "Qualitative methods",
        "Coding, reflexivity, or triangulation is reported",
        (
            r"\bthematic analysis\b",
            r"\bcoding\b",
            r"\bcodebook\b",
            r"\breflexivity\b",
            r"\btriangulation\b",
            r"\bmember checking\b",
        ),
        "Report coding procedures and researcher reflexivity checks.",
        applies_to=("qualitative",),
    ),
    Criterion(
        "Model or algorithm studies",
        "Validation split or external validation is reported",
        (
            r"\btraining set\b",
            r"\btest set\b",
            r"\bvalidation set\b",
            r"\bcross[- ]validation\b",
            r"\bexternal validation\b",
            r"\bholdout\b",
        ),
        "Describe training, validation, testing, and external validation if available.",
        applies_to=("ml",),
    ),
    Criterion(
        "Model or algorithm studies",
        "Fairness or subgroup performance is reported",
        (
            r"\bfairness\b",
            r"\bbias audit\b",
            r"\bsubgroup performance\b",
            r"\bdemographic parity\b",
            r"\bequalized odds\b",
        ),
        "Report subgroup performance and fairness checks for affected populations.",
        applies_to=("ml",),
    ),
)


BENCHMARKS: tuple[Benchmark, ...] = (
    Benchmark(
        "Protocol readiness",
        "Design, protocol, outcomes, and sample-size planning are clear before strong claims are made.",
        (
            "Study design is named",
            "Protocol or preregistration is reported",
            "Primary outcomes or endpoints are specified",
            "Sample size or power rationale is reported",
        ),
    ),
    Benchmark(
        "Participant selection fairness",
        "Recruitment, eligibility, and participant characteristics are transparent enough to judge selection bias and generalizability.",
        (
            "Eligibility criteria are described",
            "Recruitment source or study period is described",
            "Participant characteristics are reported",
        ),
    ),
    Benchmark(
        "Bias control readiness",
        "The project includes design or analysis safeguards that reduce expected sources of bias.",
        (
            "Random allocation is reported",
            "Blinding or masking is reported",
            "Comparator or control condition is described",
            "Confounding is addressed",
            "Measurement validity or reliability is addressed",
        ),
    ),
    Benchmark(
        "Ethical conduct readiness",
        "Oversight, consent, privacy, safety, or welfare safeguards are described before data collection or publication.",
        (
            "Ethics approval or oversight is reported",
            "Consent process is reported",
            "Privacy, confidentiality, or safety safeguards are reported",
            "Animal welfare oversight is reported",
        ),
    ),
    Benchmark(
        "Analysis integrity",
        "The analysis plan reports methods, uncertainty, robustness checks, and missing-data handling.",
        (
            "Statistical methods are named",
            "Uncertainty is reported",
            "Sensitivity, subgroup, or robustness checks are reported",
            "Missing data, attrition, or follow-up is addressed",
        ),
    ),
    Benchmark(
        "Transparency and conflicts",
        "Data, code, funding, sponsor role, competing interests, and limitations are visible to reviewers.",
        (
            "Data or code availability is reported",
            "Funding source is reported",
            "Conflicts or competing interests are reported",
            "Funder or sponsor role is clarified",
            "Limitations are discussed",
        ),
    ),
    Benchmark(
        "Review-method completeness",
        "Search strategy, included-study appraisal, and publication-bias assessment are planned or reported.",
        (
            "Search strategy and databases are reported",
            "Risk-of-bias or quality appraisal is reported",
            "Publication bias is considered",
        ),
        applies_to=("review",),
    ),
    Benchmark(
        "Qualitative rigor",
        "Qualitative coding, reflexivity, triangulation, or member-checking safeguards are planned or reported.",
        ("Coding, reflexivity, or triangulation is reported",),
        applies_to=("qualitative",),
    ),
    Benchmark(
        "Algorithmic fairness",
        "Model validation and subgroup or fairness checks are planned or reported for algorithmic studies.",
        (
            "Validation split or external validation is reported",
            "Fairness or subgroup performance is reported",
        ),
        applies_to=("ml",),
    ),
)


APP_CSS = """
/* ===================== OpenStudy design system ===================== */
:root {
  --os-paper: #f6f4ee;
  --os-card: #fffdf8;
  --os-ink: #20312d;
  --os-ink-soft: #5b6b66;
  --os-brand: #0e6b5f;
  --os-brand-deep: #0a554b;
  --os-brand-tint: #e3efec;
  --os-border: #e4dfd3;
  --os-serif: "Source Serif 4", "Iowan Old Style", Georgia, serif;
  --os-mono: "IBM Plex Mono", ui-monospace, SFMono-Regular, monospace;
}
.dark {
  --os-paper: #101614;
  --os-card: #19221f;
  --os-ink: #e9e7df;
  --os-ink-soft: #9fafa9;
  --os-brand: #2fa08f;
  --os-brand-deep: #3db2a0;
  --os-brand-tint: #16302b;
  --os-border: #2b3733;
}

.gradio-container {
  max-width: 1240px !important;
  margin: 0 auto !important;
}

/* ---------- Hero header ---------- */
.os-hero {
  padding: 2.1rem 0.2rem 1.5rem;
  border-bottom: 1px solid var(--os-border);
  margin-bottom: 0.35rem;
}
.os-wordmark {
  display: flex;
  align-items: center;
  gap: 0.55rem;
  font-weight: 700;
  font-size: 1.02rem;
  letter-spacing: 0.012em;
  color: var(--os-ink);
}
.os-logo {
  width: 32px;
  height: 32px;
  border-radius: 9px;
  background: var(--os-brand);
  color: #fff;
  display: grid;
  place-items: center;
  flex: none;
}
.dark .os-logo { color: #0d1a17; }
.os-hero h1 {
  font-family: var(--os-serif);
  font-weight: 600;
  font-size: clamp(1.65rem, 3.2vw, 2.35rem);
  line-height: 1.16;
  letter-spacing: -0.012em;
  color: var(--os-ink);
  max-width: 30ch;
  margin: 1rem 0 0.55rem;
}
.os-lede {
  color: var(--os-ink-soft);
  font-size: 1.02rem;
  line-height: 1.55;
  max-width: 72ch;
  margin: 0 0 1.05rem;
}
.os-badges {
  display: flex;
  flex-wrap: wrap;
  gap: 0.45rem;
  margin-bottom: 1.1rem;
}
.os-badge {
  display: inline-flex;
  align-items: center;
  gap: 0.35rem;
  padding: 0.28rem 0.72rem;
  border: 1px solid var(--os-border);
  border-radius: 999px;
  background: var(--os-card);
  color: var(--os-ink-soft);
  font-size: 0.78rem;
  font-weight: 600;
  letter-spacing: 0.01em;
  white-space: nowrap;
}
.os-badge.os-badge-brand {
  border-color: transparent;
  background: var(--os-brand-tint);
  color: var(--os-brand-deep);
}

/* ---------- Callout note ---------- */
.os-note, .openstudy-note {
  border: 1px solid var(--os-border);
  border-left: 4px solid var(--os-brand);
  padding: 0.85rem 1.05rem;
  background: var(--os-card);
  color: var(--os-ink);
  border-radius: 10px;
  font-size: 0.92rem;
  line-height: 1.55;
  max-width: 86ch;
}
.os-note strong { color: var(--os-brand-deep); }

/* ---------- Tabs as underlined nav ---------- */
.gradio-container button[role="tab"] {
  font-weight: 600;
  font-size: 0.95rem;
  color: var(--os-ink-soft);
  background: transparent;
  border-radius: 6px 6px 0 0;
  padding: 0.6rem 1rem;
  border-bottom: 2px solid transparent;
}
.gradio-container button[role="tab"][aria-selected="true"] {
  color: var(--os-brand-deep);
  border-bottom: 2px solid var(--os-brand);
  background: transparent;
}
.gradio-container .tabitem {
  padding-top: 0.85rem;
  border: none;
  background: transparent;
}

/* ---------- Typography in markdown blocks ---------- */
.gradio-container .prose h1,
.gradio-container .prose h2,
.gradio-container .prose h3 {
  font-family: var(--os-serif);
  font-weight: 600;
  letter-spacing: -0.01em;
  color: var(--os-ink);
}
.gradio-container .prose code {
  font-family: var(--os-mono);
  font-size: 0.85em;
  background: var(--os-brand-tint);
  color: var(--os-brand-deep);
  border-radius: 4px;
  padding: 0.08em 0.35em;
}

/* ---------- Buttons ---------- */
.gradio-container button.primary,
.gradio-container button.secondary {
  font-weight: 600;
  letter-spacing: 0.01em;
  transition: transform 0.08s ease, box-shadow 0.12s ease;
}
.gradio-container button.primary:hover {
  box-shadow: 0 3px 10px rgba(14, 107, 95, 0.28);
}
.gradio-container button.primary:active,
.gradio-container button.secondary:active {
  transform: translateY(1px);
}

/* ---------- Tables ---------- */
.gradio-container .table-wrap {
  border: 1px solid var(--os-border);
  border-radius: 10px;
}
.gradio-container thead th,
.gradio-container thead th span {
  font-size: 0.72rem !important;
  text-transform: uppercase;
  letter-spacing: 0.07em;
  font-weight: 700;
  color: var(--os-ink-soft);
}
.gradio-container tbody td,
.gradio-container tbody td span {
  font-size: 0.86rem;
  line-height: 1.45;
}
.compact-table textarea { font-size: 0.88rem !important; }

/* ---------- Inputs ---------- */
.gradio-container input:focus,
.gradio-container textarea:focus {
  border-color: var(--os-brand) !important;
  box-shadow: 0 0 0 3px rgba(14, 107, 95, 0.14) !important;
}

/* ---------- Footer ---------- */
.os-footer {
  display: flex;
  flex-wrap: wrap;
  justify-content: space-between;
  gap: 0.5rem 1.5rem;
  border-top: 1px solid var(--os-border);
  margin-top: 2.2rem;
  padding: 1.1rem 0.2rem 0.4rem;
  color: var(--os-ink-soft);
  font-size: 0.82rem;
}
.os-footer a { color: var(--os-brand); text-decoration: none; }

@media (max-width: 720px) {
  .os-hero { padding-top: 1.4rem; }
  .os-badge { font-size: 0.72rem; }
}
"""


def strip_markup(value: Any) -> str:
    if not value:
        return ""
    text = str(value)
    text = re.sub(r"<[^>]+>", " ", text)
    text = html.unescape(text)
    return normalize_space(text)


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def normalize_doi(value: str) -> str:
    match = re.search(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b", value or "", flags=re.I)
    if not match:
        return ""
    return match.group(0).rstrip(".,;)]}").lower()


def is_url(value: str) -> bool:
    return bool(re.match(r"^https?://", (value or "").strip(), flags=re.I))


def safe_get_json(url: str, params: dict[str, Any] | None = None) -> dict[str, Any]:
    try:
        response = requests.get(url, params=params, headers=HEADERS, timeout=REQUEST_TIMEOUT)
        response.raise_for_status()
        return response.json()
    except Exception:
        return {}


def get_first(value: Any) -> str:
    if isinstance(value, list) and value:
        return strip_markup(value[0])
    return strip_markup(value)


def year_from_crossref(item: dict[str, Any]) -> str:
    for key in ("published-print", "published-online", "published", "issued", "created"):
        parts = item.get(key, {}).get("date-parts", [])
        if parts and parts[0]:
            return str(parts[0][0])
    return ""


def record_from_crossref(item: dict[str, Any]) -> dict[str, Any]:
    title = get_first(item.get("title"))
    container = get_first(item.get("container-title"))
    abstract = strip_markup(item.get("abstract"))
    doi = strip_markup(item.get("DOI")).lower()
    authors = []
    for author in item.get("author", [])[:8]:
        name = normalize_space(" ".join(part for part in (author.get("given"), author.get("family")) if part))
        if name:
            authors.append(name)
    return {
        "title": title or "Untitled study",
        "year": year_from_crossref(item),
        "venue": container,
        "doi": doi,
        "url": item.get("URL", ""),
        "abstract": abstract,
        "authors": ", ".join(authors),
        "source": "Crossref",
    }


def crossref_by_doi(doi: str) -> dict[str, Any] | None:
    url = f"https://api.crossref.org/works/{quote(doi, safe='')}"
    data = safe_get_json(url)
    item = data.get("message")
    if isinstance(item, dict):
        return record_from_crossref(item)
    return None


def crossref_search(query: str, rows: int) -> list[dict[str, Any]]:
    data = safe_get_json(
        "https://api.crossref.org/works",
        params={"query.bibliographic": query, "rows": rows, "select": "title,DOI,URL,abstract,author,container-title,published,published-print,published-online,issued,created"},
    )
    items = data.get("message", {}).get("items", [])
    return [record_from_crossref(item) for item in items if isinstance(item, dict)]


def abstract_from_openalex(index: dict[str, list[int]] | None) -> str:
    if not index:
        return ""
    positions: dict[int, str] = {}
    for word, offsets in index.items():
        for offset in offsets:
            positions[int(offset)] = word
    return normalize_space(" ".join(positions[i] for i in sorted(positions)))


def record_from_openalex(item: dict[str, Any]) -> dict[str, Any]:
    primary = item.get("primary_location") or {}
    source = primary.get("source") or {}
    doi = (item.get("doi") or "").replace("https://doi.org/", "").lower()
    authors = []
    for authorship in item.get("authorships", [])[:8]:
        author = authorship.get("author", {})
        name = strip_markup(author.get("display_name"))
        if name:
            authors.append(name)
    return {
        "title": strip_markup(item.get("display_name")) or "Untitled study",
        "year": str(item.get("publication_year") or ""),
        "venue": strip_markup(source.get("display_name")),
        "doi": doi,
        "url": item.get("doi") or item.get("id") or "",
        "abstract": abstract_from_openalex(item.get("abstract_inverted_index")),
        "authors": ", ".join(authors),
        "source": "OpenAlex",
    }


def openalex_by_doi(doi: str) -> dict[str, Any] | None:
    url = f"https://api.openalex.org/works/https://doi.org/{doi}"
    data = safe_get_json(url)
    if data.get("id"):
        return record_from_openalex(data)
    return None


def openalex_search(query: str, rows: int) -> list[dict[str, Any]]:
    data = safe_get_json("https://api.openalex.org/works", params={"search": query, "per-page": rows})
    items = data.get("results", [])
    return [record_from_openalex(item) for item in items if isinstance(item, dict)]


def extract_page_record(url: str) -> dict[str, Any] | None:
    try:
        response = requests.get(url, headers=HEADERS, timeout=REQUEST_TIMEOUT)
        response.raise_for_status()
    except Exception:
        return None

    content_type = response.headers.get("content-type", "").lower()
    if "pdf" in content_type or url.lower().endswith(".pdf"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(response.content)
            tmp_path = tmp.name
        text = extract_text_from_file(tmp_path)
        return {
            "title": Path(url).name or "Uploaded PDF from URL",
            "year": "",
            "venue": "",
            "doi": normalize_doi(text),
            "url": url,
            "abstract": text[:6000],
            "authors": "",
            "source": "PDF URL",
        }

    soup = BeautifulSoup(response.text, "html.parser")

    def meta_content(*names: str) -> str:
        for name in names:
            tag = soup.find("meta", attrs={"name": name}) or soup.find("meta", attrs={"property": name})
            if tag and tag.get("content"):
                return strip_markup(tag.get("content"))
        return ""

    title = meta_content("citation_title", "dc.title", "og:title") or strip_markup(soup.title.string if soup.title else "")
    abstract = meta_content("citation_abstract", "dc.description", "description", "og:description")
    doi = meta_content("citation_doi", "dc.identifier") or normalize_doi(response.text)
    year = meta_content("citation_publication_date", "citation_online_date")[:4]
    venue = meta_content("citation_journal_title", "citation_conference_title")
    if doi:
        richer = crossref_by_doi(doi) or openalex_by_doi(doi)
        if richer:
            if not richer.get("abstract") and abstract:
                richer["abstract"] = abstract
            richer["source"] = f"{richer['source']} + page"
            return richer
    return {
        "title": title or url,
        "year": year,
        "venue": venue,
        "doi": doi.lower(),
        "url": url,
        "abstract": abstract,
        "authors": "",
        "source": "Web page",
    }


def merge_records(records: list[dict[str, Any]], limit: int) -> list[dict[str, Any]]:
    merged: dict[str, dict[str, Any]] = {}
    for record in records:
        if not record:
            continue
        doi = normalize_doi(record.get("doi", ""))
        title_key = re.sub(r"[^a-z0-9]+", "", (record.get("title") or "").lower())[:90]
        key = doi or title_key
        if not key:
            continue
        if key not in merged:
            merged[key] = record
            continue
        existing = merged[key]
        for field in ("title", "year", "venue", "doi", "url", "authors"):
            if not existing.get(field) and record.get(field):
                existing[field] = record[field]
        if len(record.get("abstract", "")) > len(existing.get("abstract", "")):
            existing["abstract"] = record["abstract"]
        sources = {part.strip() for part in existing.get("source", "").split("+") if part.strip()}
        sources.update(part.strip() for part in record.get("source", "").split("+") if part.strip())
        existing["source"] = " + ".join(sorted(sources))
    return list(merged.values())[:limit]


def search_studies(query: str, rows: int) -> list[dict[str, Any]]:
    query = normalize_space(query)
    if not query:
        return []
    rows = max(1, min(int(rows), 10))
    records: list[dict[str, Any]] = []
    doi = normalize_doi(query)

    if doi:
        records.extend(record for record in (crossref_by_doi(doi), openalex_by_doi(doi)) if record)
        if records:
            return merge_records(records, rows)

    if is_url(query) and not records:
        page_record = extract_page_record(query)
        if page_record:
            records.append(page_record)
            return merge_records(records, rows)

    if not records or len(records) < rows:
        records.extend(crossref_search(query, rows))
        records.extend(openalex_search(query, rows))

    return merge_records(records, rows)


def detect_study_type(text: str) -> tuple[str, str]:
    lowered = text.lower()
    scores = {
        "review": count_patterns(lowered, (r"\bsystematic review\b", r"\bmeta[- ]analysis\b", r"\bprisma\b")),
        "clinical_trial": count_patterns(lowered, (r"\brandomi[sz]ed controlled trial\b", r"\bclinical trial\b", r"\btrial registration\b", r"\bplacebo\b")),
        "observational": count_patterns(lowered, (r"\bcohort\b", r"\bcase[- ]control\b", r"\bcross[- ]sectional\b", r"\bretrospective\b", r"\bprospective observational\b")),
        "qualitative": count_patterns(lowered, (r"\bqualitative\b", r"\binterviews?\b", r"\bfocus groups?\b", r"\bthematic analysis\b")),
        "animal": count_patterns(lowered, (r"\bmice\b", r"\brats?\b", r"\banimal model\b", r"\biacuc\b", r"\banimal care\b")),
        "ml": count_patterns(lowered, (r"\bmachine learning\b", r"\bdeep learning\b", r"\bprediction model\b", r"\balgorithm\b", r"\bclassifier\b")),
        "survey": count_patterns(lowered, (r"\bsurvey\b", r"\bquestionnaire\b", r"\brespondents\b")),
    }
    best_key, best_score = max(scores.items(), key=lambda item: item[1])
    labels = {
        "review": "Systematic review or meta-analysis",
        "clinical_trial": "Clinical trial",
        "observational": "Observational study",
        "qualitative": "Qualitative study",
        "animal": "Animal or preclinical study",
        "ml": "Model or algorithm study",
        "survey": "Survey study",
        "unknown": "Not enough text to classify",
    }
    if best_score <= 0:
        return "unknown", labels["unknown"]
    return best_key, labels[best_key]


def count_patterns(text: str, patterns: tuple[str, ...]) -> int:
    return sum(1 for pattern in patterns if re.search(pattern, text, flags=re.I))


def criterion_applies(criterion: Criterion, study_key: str) -> bool:
    if study_key in criterion.excludes:
        return False
    if criterion.applies_to:
        return study_key in criterion.applies_to
    return True


def benchmark_applies(benchmark: Benchmark, study_key: str) -> bool:
    if study_key in benchmark.excludes:
        return False
    if benchmark.applies_to:
        return study_key in benchmark.applies_to
    return True


def split_sentences(text: str) -> list[str]:
    normalized = normalize_space(text)
    pieces = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", normalized)
    if len(pieces) <= 1:
        pieces = re.split(r"\s{2,}", normalized)
    return [piece.strip() for piece in pieces if piece.strip()]


def find_snippet(text: str, patterns: tuple[str, ...], max_chars: int = 260) -> str:
    sentences = split_sentences(text[:MAX_ANALYSIS_CHARS])
    for sentence in sentences:
        for pattern in patterns:
            if re.search(pattern, sentence, flags=re.I):
                return truncate(sentence, max_chars)
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.I)
        if match:
            start = max(0, match.start() - max_chars // 2)
            end = min(len(text), match.end() + max_chars // 2)
            return truncate(normalize_space(text[start:end]), max_chars)
    return ""


def truncate(text: str, max_chars: int) -> str:
    text = normalize_space(text)
    if len(text) <= max_chars:
        return text
    return text[: max_chars - 3].rstrip() + "..."


def evaluate_criterion(text: str, criterion: Criterion) -> dict[str, Any]:
    positive_hit = any(re.search(pattern, text, flags=re.I) for pattern in criterion.positive)
    concern_hit = any(re.search(pattern, text, flags=re.I) for pattern in criterion.concern)
    if positive_hit:
        status = "Reported"
        score = 1.0
        evidence = find_snippet(text, criterion.positive)
    elif concern_hit:
        status = "Potential concern"
        score = 0.25
        evidence = find_snippet(text, criterion.concern)
    else:
        status = "Missing or unclear"
        score = 0.0
        evidence = ""
    return {
        "category": criterion.category,
        "criterion": criterion.label,
        "status": status,
        "score": score,
        "evidence": evidence,
        "recommendation": criterion.recommendation,
    }


def category_scores(results: list[dict[str, Any]]) -> dict[str, float]:
    grouped: dict[str, list[float]] = {}
    for result in results:
        grouped.setdefault(result["category"], []).append(float(result["score"]))
    return {category: round(sum(values) / len(values) * 100, 1) for category, values in grouped.items()}


def risk_level(score: float) -> tuple[str, str]:
    if score >= 75:
        return "Lower reporting risk", "The available text reports many safeguards readers expect."
    if score >= 50:
        return "Moderate reporting risk", "Several safeguards are reported, but important items need review."
    return "Higher reporting risk", "Many expected safeguards are missing or unclear in the available text."


def evaluate_study_text(
    text: str,
    metadata: dict[str, Any] | None = None,
    scope: str = "available text",
    forced_study_key: str | None = None,
) -> dict[str, Any]:
    metadata = metadata or {}
    text = normalize_space(text)
    truncated = len(text) > MAX_ANALYSIS_CHARS
    analysis_text = text[:MAX_ANALYSIS_CHARS]
    detected_key, detected_label = detect_study_type(analysis_text)
    study_key = forced_study_key or detected_key
    study_label = label_for_study_key(study_key) if forced_study_key else detected_label
    criteria = [criterion for criterion in CRITERIA if criterion_applies(criterion, study_key)]
    checklist = [evaluate_criterion(analysis_text, criterion) for criterion in criteria]
    scores = category_scores(checklist)
    overall = round(sum(item["score"] for item in checklist) / max(len(checklist), 1) * 100, 1)
    level, level_detail = risk_level(overall)
    conducted = extract_conduct_summary(analysis_text, metadata, study_label)
    benchmarks = build_benchmark_rows(checklist, study_key)
    biases = build_bias_rows(checklist, study_key)
    return {
        "metadata": metadata,
        "scope": scope,
        "study_key": study_key,
        "study_label": study_label,
        "checklist": checklist,
        "scores": scores,
        "overall": overall,
        "level": level,
        "level_detail": level_detail,
        "conducted": conducted,
        "benchmarks": benchmarks,
        "biases": biases,
        "truncated": truncated,
        "text_chars": len(text),
    }


def label_for_study_key(study_key: str) -> str:
    labels = {
        "review": "Systematic review or meta-analysis",
        "clinical_trial": "Clinical trial",
        "observational": "Observational study",
        "qualitative": "Qualitative study",
        "animal": "Animal or preclinical study",
        "ml": "Model or algorithm study",
        "survey": "Survey study",
        "unknown": "Not enough text to classify",
    }
    return labels.get(study_key, labels["unknown"])


def key_for_type_hint(type_hint: str) -> str | None:
    mapping = {
        "Clinical trial": "clinical_trial",
        "Observational study": "observational",
        "Systematic review or meta-analysis": "review",
        "Qualitative study": "qualitative",
        "Survey study": "survey",
        "Animal or preclinical study": "animal",
        "Model or algorithm study": "ml",
    }
    return mapping.get(type_hint or "")


def extract_conduct_summary(text: str, metadata: dict[str, Any], study_label: str) -> list[tuple[str, str]]:
    fields = [
        ("Title", metadata.get("title") or find_snippet(text, (r"^.{20,180}",), 180)),
        ("Detected study type", study_label),
    ]
    if metadata.get("project_status"):
        fields.append(("Project status", metadata["project_status"]))
    if metadata.get("team"):
        fields.append(("Team or owner", metadata["team"]))
    fields.extend(
        [
            ("Authors", metadata.get("authors", "")),
            ("Publication", " ".join(part for part in (metadata.get("venue"), metadata.get("year")) if part)),
            ("DOI or URL", metadata.get("doi") or metadata.get("url", "")),
            ("Population or sample", find_snippet(text, (r"\bparticipants?\b", r"\bpatients?\b", r"\bsample\b", r"\bcohort\b", r"\brecruited\b", r"\benrolled\b"))),
            ("Intervention, exposure, or focus", find_snippet(text, (r"\bintervention\b", r"\btreatment\b", r"\bexposure\b", r"\bprogram\b", r"\bassigned\b", r"\bassociation between\b"))),
            ("Comparator or control", find_snippet(text, (r"\bcontrol group\b", r"\bcomparator\b", r"\bplacebo\b", r"\busual care\b", r"\bmatched controls?\b"))),
            ("Outcomes or endpoints", find_snippet(text, (r"\bprimary outcome\b", r"\bendpoint\b", r"\boutcome measures?\b", r"\bmeasured\b"))),
            ("Analysis approach", find_snippet(text, (r"\bregression\b", r"\banova\b", r"\bodds ratio\b", r"\bhazard ratio\b", r"\bconfidence interval\b", r"\bmodel\b"))),
            ("Ethics or consent", find_snippet(text, (r"\bethics\b", r"\birb\b", r"\binformed consent\b", r"\bconsent\b", r"\bdeclaration of helsinki\b"))),
            ("Funding or conflicts", find_snippet(text, (r"\bfunding\b", r"\bgrant\b", r"\bsponsor\b", r"\bconflict", r"\bcompeting interests?\b"))),
        ]
    )
    return [(label, value or "Not found in available text") for label, value in fields]


def checklist_lookup(checklist: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    return {item["criterion"]: item for item in checklist}


def combine_evidence(items: list[dict[str, Any]]) -> str:
    snippets = [item["evidence"] for item in items if item.get("evidence")]
    if snippets:
        return " | ".join(snippets[:2])
    missing = [item["criterion"] for item in items if item.get("status") == "Missing or unclear"]
    if missing:
        return "Missing or unclear: " + "; ".join(missing[:3])
    return "No strong concern found in available text"


def bias_status(items: list[dict[str, Any]]) -> str:
    statuses = [item.get("status") for item in items]
    if any(status == "Potential concern" for status in statuses):
        return "Potential concern"
    if all(status == "Reported" for status in statuses if status):
        return "Lower signal"
    return "Needs review"


def benchmark_status(score: float) -> str:
    if score >= 0.85:
        return "Meets benchmark"
    if score >= 0.55:
        return "Partially meets benchmark"
    return "Needs work"


def build_benchmark_rows(checklist: list[dict[str, Any]], study_key: str) -> list[list[str]]:
    lookup = checklist_lookup(checklist)
    rows = []
    for benchmark in BENCHMARKS:
        if not benchmark_applies(benchmark, study_key):
            continue
        items = [lookup[label] for label in benchmark.criteria if label in lookup]
        if not items:
            continue
        score = sum(float(item["score"]) for item in items) / len(items)
        missing = [item["criterion"] for item in items if item["status"] != "Reported"]
        action = "Maintain this benchmark and keep evidence in the protocol or report."
        if missing:
            action = "Address: " + "; ".join(missing[:4])
        rows.append(
            [
                benchmark.name,
                benchmark_status(score),
                f"{score * 100:.0f}%",
                benchmark.description,
                combine_evidence(items),
                action,
            ]
        )
    return rows


def build_bias_rows(checklist: list[dict[str, Any]], study_key: str) -> list[list[str]]:
    lookup = checklist_lookup(checklist)

    def pick(*labels: str) -> list[dict[str, Any]]:
        return [lookup[label] for label in labels if label in lookup]

    rows = [
        [
            "Selection bias",
            bias_status(pick("Eligibility criteria are described", "Recruitment source or study period is described", "Participant characteristics are reported")),
            combine_evidence(pick("Eligibility criteria are described", "Recruitment source or study period is described", "Participant characteristics are reported")),
            "Review whether the sample represents the target population and whether exclusions could skew results.",
        ],
        [
            "Reporting bias",
            bias_status(pick("Protocol or preregistration is reported", "Primary outcomes or endpoints are specified", "Limitations are discussed")),
            combine_evidence(pick("Protocol or preregistration is reported", "Primary outcomes or endpoints are specified", "Limitations are discussed")),
            "Check for preregistration, outcome switching, selective reporting, and absent limitations.",
        ],
        [
            "Funding or conflict bias",
            bias_status(pick("Funding source is reported", "Conflicts or competing interests are reported", "Funder or sponsor role is clarified")),
            combine_evidence(pick("Funding source is reported", "Conflicts or competing interests are reported", "Funder or sponsor role is clarified")),
            "Check whether funders or sponsors influenced design, analysis, interpretation, or publication.",
        ],
        [
            "Ethics safeguards",
            bias_status(pick("Ethics approval or oversight is reported", "Consent process is reported", "Privacy, confidentiality, or safety safeguards are reported", "Animal welfare oversight is reported")),
            combine_evidence(pick("Ethics approval or oversight is reported", "Consent process is reported", "Privacy, confidentiality, or safety safeguards are reported", "Animal welfare oversight is reported")),
            "Confirm oversight, consent or waiver, privacy protections, and participant or animal welfare safeguards.",
        ],
        [
            "Measurement bias",
            bias_status(pick("Measurement validity or reliability is addressed", "Blinding or masking is reported")),
            combine_evidence(pick("Measurement validity or reliability is addressed", "Blinding or masking is reported")),
            "Check whether outcomes were measured with valid instruments and protected from assessor expectations.",
        ],
        [
            "Attrition or missing-data bias",
            bias_status(pick("Missing data, attrition, or follow-up is addressed")),
            combine_evidence(pick("Missing data, attrition, or follow-up is addressed")),
            "Look for lost-to-follow-up counts, exclusions after enrollment, imputation, and sensitivity checks.",
        ],
    ]

    if study_key in {"observational", "survey", "unknown"}:
        rows.append(
            [
                "Confounding",
                bias_status(pick("Confounding is addressed", "Causal claims are appropriately constrained")),
                combine_evidence(pick("Confounding is addressed", "Causal claims are appropriately constrained")),
                "For non-randomized designs, check whether major confounders were identified and adjusted.",
            ]
        )
    if study_key == "clinical_trial":
        rows.append(
            [
                "Performance and detection bias",
                bias_status(pick("Random allocation is reported", "Blinding or masking is reported", "Comparator or control condition is described")),
                combine_evidence(pick("Random allocation is reported", "Blinding or masking is reported", "Comparator or control condition is described")),
                "Confirm allocation concealment, blinding, and comparator details.",
            ]
        )
    if study_key == "review":
        rows.append(
            [
                "Review publication bias",
                bias_status(pick("Search strategy and databases are reported", "Risk-of-bias or quality appraisal is reported", "Publication bias is considered")),
                combine_evidence(pick("Search strategy and databases are reported", "Risk-of-bias or quality appraisal is reported", "Publication bias is considered")),
                "Review search completeness, included-study appraisal, and publication-bias assessment.",
            ]
        )
    if study_key == "ml":
        rows.append(
            [
                "Algorithmic bias",
                bias_status(pick("Validation split or external validation is reported", "Fairness or subgroup performance is reported", "Participant characteristics are reported")),
                combine_evidence(pick("Validation split or external validation is reported", "Fairness or subgroup performance is reported", "Participant characteristics are reported")),
                "Check subgroup performance, external validation, leakage, and whether affected groups are represented.",
            ]
        )
    return rows


def style_plot(fig, ax) -> None:
    fig.patch.set_facecolor(CARD)
    ax.set_facecolor(CARD)
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.tick_params(colors=INK, length=0, labelsize=9.5)


def score_plot(scores: dict[str, float]):
    fig, ax = plt.subplots(figsize=(8.5, max(3.2, 0.42 * len(scores))))
    style_plot(fig, ax)
    if not scores:
        ax.text(0.5, 0.5, "No scores available", ha="center", va="center", color=INK_SOFT)
        ax.axis("off")
        return fig

    ordered = sorted(scores.items(), key=lambda item: item[1])
    labels = [item[0] for item in ordered]
    values = [item[1] for item in ordered]
    colors = [RISK if value < 50 else WARN if value < 75 else BRAND for value in values]
    ax.barh(labels, values, color=colors, height=0.62, zorder=3)
    ax.barh(labels, [100] * len(values), color=BORDER, height=0.62, zorder=1, alpha=0.45)
    ax.set_xlim(0, 100)
    ax.set_xlabel("Reported safeguard score", color=INK_SOFT, fontsize=9.5)
    ax.grid(axis="x", color=BORDER, linewidth=0.8, zorder=0)
    for index, value in enumerate(values):
        inside = value > 88
        ax.text(
            value - 2 if inside else value + 1.5,
            index,
            f"{value:.0f}%",
            va="center",
            ha="right" if inside else "left",
            fontsize=9,
            fontweight="bold",
            color="#ffffff" if inside else INK,
            zorder=4,
        )
    ax.set_title("Reporting safeguards by category", loc="left", color=INK, fontsize=11, fontweight="bold", pad=12)
    fig.tight_layout()
    return fig


def summary_markdown(report: dict[str, Any]) -> str:
    metadata = report["metadata"]
    title = metadata.get("title") or "Untitled study"
    caution = " This upload was truncated for analysis." if report["truncated"] else ""
    return f"""
### {strip_markup(title)}

**Overall screen:** {report['overall']:.1f}/100 - **{report['level']}**

{report['level_detail']} This is a transparent screening of the **{report['scope']}**, not proof that the research is ethical, unethical, biased, or unbiased.{caution}

**Detected study type:** {report['study_label']}  
**Text reviewed:** {report['text_chars']:,} characters
"""


def conduct_markdown(report: dict[str, Any]) -> str:
    rows = report["conducted"]
    body = "\n".join(f"- **{label}:** {value}" for label, value in rows if value)
    return "### How the study appears to have been conducted\n" + body


def checklist_rows(report: dict[str, Any]) -> list[list[str]]:
    return [
        [
            item["category"],
            item["criterion"],
            item["status"],
            item["evidence"] or "Not found in available text",
            item["recommendation"],
        ]
        for item in report["checklist"]
    ]


def report_to_markdown(report: dict[str, Any]) -> str:
    lines = [
        f"# {APP_NAME} screening report",
        "",
        summary_markdown(report).strip(),
        "",
        conduct_markdown(report).strip(),
        "",
        "## Benchmarks",
        "",
        "| Benchmark | Status | Score | Description | Evidence | Required action |",
        "| --- | --- | --- | --- | --- | --- |",
    ]
    for benchmark, status, score, description, evidence, action in report.get("benchmarks", []):
        lines.append(
            f"| {escape_table(benchmark)} | {escape_table(status)} | {escape_table(score)} | {escape_table(description)} | {escape_table(evidence)} | {escape_table(action)} |"
        )
    lines.extend(
        [
            "",
        "## Potential bias and ethics signals",
        "",
        "| Signal | Status | Evidence | Next check |",
        "| --- | --- | --- | --- |",
        ]
    )
    for signal, status, evidence, next_check in report["biases"]:
        lines.append(f"| {escape_table(signal)} | {escape_table(status)} | {escape_table(evidence)} | {escape_table(next_check)} |")
    lines.extend(
        [
            "",
            "## Checklist",
            "",
            "| Category | Criterion | Status | Evidence | Recommendation |",
            "| --- | --- | --- | --- | --- |",
        ]
    )
    for row in checklist_rows(report):
        lines.append("| " + " | ".join(escape_table(value) for value in row) + " |")
    lines.extend(
        [
            "",
            "## Important limitation",
            "",
            "This report screens reported safeguards in the text supplied to the app. It does not replace peer review, domain expert review, IRB review, legal advice, or replication.",
        ]
    )
    return "\n".join(lines)


def escape_table(value: Any) -> str:
    return normalize_space(str(value)).replace("|", "\\|")


def write_report_file(report: dict[str, Any]) -> str:
    with tempfile.NamedTemporaryFile("w", delete=False, suffix=".md", encoding="utf-8") as handle:
        handle.write(report_to_markdown(report))
        return handle.name


def format_choice(index: int, record: dict[str, Any]) -> str:
    year = f" ({record.get('year')})" if record.get("year") else ""
    title = truncate(record.get("title") or "Untitled study", 88)
    return f"{index + 1}. {title}{year}"


def search_studies_ui(query: str, rows: int):
    records = search_studies(query, int(rows or 5))
    table = []
    for index, record in enumerate(records, start=1):
        table.append(
            [
                index,
                record.get("title", ""),
                record.get("year", ""),
                record.get("venue", ""),
                record.get("doi", ""),
                record.get("source", ""),
                truncate(record.get("abstract", ""), 260) or "No abstract found",
            ]
        )
    choices = [format_choice(index, record) for index, record in enumerate(records)]
    status = f"Found {len(records)} candidate studies. Select one to evaluate." if records else "No candidate studies found. Try a DOI, a more exact title, or paste an abstract in the Upload tab."
    return table, gr.update(choices=choices, value=choices[0] if choices else None), records, status


def analyze_selected_ui(choice: str, records: list[dict[str, Any]] | None):
    records = records or []
    if not choice or not records:
        return empty_outputs("Search for studies first, then select a result.")
    match = re.match(r"^(\d+)\.", choice)
    index = int(match.group(1)) - 1 if match else 0
    if index < 0 or index >= len(records):
        return empty_outputs("Selected study was not found in the current search results.")
    record = records[index]
    text_parts = [record.get("title", ""), record.get("abstract", "")]
    if not normalize_space(" ".join(text_parts)):
        return empty_outputs("That result does not include enough text to evaluate. Paste the abstract or upload the manuscript.")
    report = evaluate_study_text(" ".join(text_parts), record, scope="metadata and abstract available from public search")
    return report_outputs(report)


def uploaded_file_text(file_path: Any) -> tuple[str, str]:
    if file_path:
        path = file_path
        if isinstance(file_path, dict):
            path = file_path.get("path") or file_path.get("name")
        elif hasattr(file_path, "name"):
            path = file_path.name
        return Path(str(path)).name, extract_text_from_file(str(path))
    return "pasted text", ""


def analyze_upload_ui(file_path: Any, pasted_text: str, title: str, study_type_hint: str):
    source_name, extracted = uploaded_file_text(file_path)
    text = normalize_space(" ".join(part for part in (title, pasted_text, extracted) if part))
    if not text:
        return empty_outputs("Upload a PDF/text file or paste manuscript text to evaluate.")
    metadata = {"title": title or source_name, "source": source_name}
    report = evaluate_study_text(
        text,
        metadata,
        scope="uploaded or pasted manuscript text",
        forced_study_key=key_for_type_hint(study_type_hint),
    )
    return report_outputs(report)


def project_text_from_fields(
    uploaded_project_text: str,
    title: str,
    team: str,
    project_status: str,
    study_type_hint: str,
    research_question: str,
    study_design: str,
    participants: str,
    recruitment: str,
    eligibility: str,
    intervention_or_exposure: str,
    comparator: str,
    outcomes: str,
    sample_size: str,
    bias_controls: str,
    ethics_plan: str,
    consent_privacy: str,
    analysis_plan: str,
    missing_data_plan: str,
    transparency_plan: str,
    funding_conflicts: str,
    notes: str,
) -> str:
    fields = [
        ("Uploaded project or protocol file text", uploaded_project_text),
        ("Project title", title),
        ("Team or owner", team),
        ("Project status", project_status),
        ("Study type", study_type_hint if study_type_hint != "Auto-detect" else ""),
        ("Research question", research_question),
        ("Study design", study_design),
        ("Population or sample", participants),
        ("Recruitment source and study period", recruitment),
        ("Inclusion and exclusion criteria", eligibility),
        ("Intervention, exposure, or focus", intervention_or_exposure),
        ("Comparator or control", comparator),
        ("Primary and secondary outcomes", outcomes),
        ("Sample size or power rationale", sample_size),
        ("Bias controls", bias_controls),
        ("Ethics approval and oversight", ethics_plan),
        ("Consent, privacy, and safety safeguards", consent_privacy),
        ("Statistical or qualitative analysis plan", analysis_plan),
        ("Missing data, attrition, or follow-up plan", missing_data_plan),
        ("Data, code, materials, and preregistration plan", transparency_plan),
        ("Funding, sponsor role, and conflicts of interest", funding_conflicts),
        ("Additional notes", notes),
    ]
    return "\n\n".join(f"{label}: {value}" for label, value in fields if normalize_space(value))


def top_project_gaps(report: dict[str, Any], limit: int = 4) -> str:
    priority_statuses = {"Potential concern", "Missing or unclear"}
    gaps = [
        item["criterion"]
        for item in report["checklist"]
        if item.get("status") in priority_statuses
    ]
    return "; ".join(gaps[:limit]) if gaps else "No major gaps flagged"


def project_table_rows(projects: list[dict[str, Any]] | None) -> list[list[Any]]:
    rows = []
    for index, project in enumerate(projects or [], start=1):
        report = project.get("report", {})
        rows.append(
            [
                index,
                project.get("title", "Untitled project"),
                report.get("study_label", project.get("study_type", "")),
                project.get("project_status", ""),
                f"{report.get('overall', 0):.1f}",
                report.get("level", ""),
                top_project_gaps(report, limit=3) if report else "",
            ]
        )
    return rows


def project_choices(projects: list[dict[str, Any]] | None) -> list[str]:
    choices = []
    for index, project in enumerate(projects or [], start=1):
        title = truncate(project.get("title") or "Untitled project", 80)
        choices.append(f"{index}. {title}")
    return choices


def make_project_record(
    file_path: Any,
    title: str,
    team: str,
    project_status: str,
    study_type_hint: str,
    research_question: str,
    study_design: str,
    participants: str,
    recruitment: str,
    eligibility: str,
    intervention_or_exposure: str,
    comparator: str,
    outcomes: str,
    sample_size: str,
    bias_controls: str,
    ethics_plan: str,
    consent_privacy: str,
    analysis_plan: str,
    missing_data_plan: str,
    transparency_plan: str,
    funding_conflicts: str,
    notes: str,
) -> dict[str, Any] | None:
    source_name, uploaded_project_text = uploaded_file_text(file_path)
    project_text = project_text_from_fields(
        uploaded_project_text,
        title,
        team,
        project_status,
        study_type_hint,
        research_question,
        study_design,
        participants,
        recruitment,
        eligibility,
        intervention_or_exposure,
        comparator,
        outcomes,
        sample_size,
        bias_controls,
        ethics_plan,
        consent_privacy,
        analysis_plan,
        missing_data_plan,
        transparency_plan,
        funding_conflicts,
        notes,
    )
    if not project_text:
        return None

    display_title = normalize_space(title) or truncate(research_question, 80) or "Untitled research project"
    metadata = {
        "title": display_title,
        "team": normalize_space(team),
        "project_status": normalize_space(project_status),
        "source": source_name if uploaded_project_text else "User research project form",
    }
    report = evaluate_study_text(
        project_text,
        metadata,
        scope="research project plan supplied by the user",
        forced_study_key=key_for_type_hint(study_type_hint),
    )
    return {
        "title": display_title,
        "team": normalize_space(team),
        "project_status": normalize_space(project_status),
        "study_type": report["study_label"],
        "source": metadata["source"],
        "created_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "text": project_text,
        "report": report,
    }


def add_project_ui(
    file_path: Any,
    title: str,
    team: str,
    project_status: str,
    study_type_hint: str,
    research_question: str,
    study_design: str,
    participants: str,
    recruitment: str,
    eligibility: str,
    intervention_or_exposure: str,
    comparator: str,
    outcomes: str,
    sample_size: str,
    bias_controls: str,
    ethics_plan: str,
    consent_privacy: str,
    analysis_plan: str,
    missing_data_plan: str,
    transparency_plan: str,
    funding_conflicts: str,
    notes: str,
    projects: list[dict[str, Any]] | None,
):
    project = make_project_record(
        file_path,
        title,
        team,
        project_status,
        study_type_hint,
        research_question,
        study_design,
        participants,
        recruitment,
        eligibility,
        intervention_or_exposure,
        comparator,
        outcomes,
        sample_size,
        bias_controls,
        ethics_plan,
        consent_privacy,
        analysis_plan,
        missing_data_plan,
        transparency_plan,
        funding_conflicts,
        notes,
    )
    if not project:
        return (
            project_table_rows(projects),
            gr.update(choices=project_choices(projects), value=None),
            projects or [],
            "Add at least a title, research question, or project plan detail.",
            *empty_outputs("Add project details first."),
        )

    updated_projects = [*(projects or []), project]
    choices = project_choices(updated_projects)
    status = f"Added {project['title']} to this session and reviewed the current plan."
    return (
        project_table_rows(updated_projects),
        gr.update(choices=choices, value=choices[-1]),
        updated_projects,
        status,
        *report_outputs(project["report"]),
    )


def review_project_ui(choice: str, projects: list[dict[str, Any]] | None):
    projects = projects or []
    if not choice or not projects:
        return empty_outputs("Add a research project first, then select it for review.")
    match = re.match(r"^(\d+)\.", choice)
    index = int(match.group(1)) - 1 if match else 0
    if index < 0 or index >= len(projects):
        return empty_outputs("Selected project was not found in this session.")
    return report_outputs(projects[index]["report"])


def download_projects_ui(projects: list[dict[str, Any]] | None):
    projects = projects or []
    if not projects:
        return None
    exportable = []
    for project in projects:
        report = project.get("report", {})
        exportable.append(
            {
                "title": project.get("title"),
                "team": project.get("team"),
                "project_status": project.get("project_status"),
                "study_type": report.get("study_label"),
                "created_at": project.get("created_at"),
                "overall_score": report.get("overall"),
                "risk_level": report.get("level"),
                "source": project.get("source"),
                "benchmarks": [
                    {
                        "benchmark": row[0],
                        "status": row[1],
                        "score": row[2],
                        "required_action": row[5],
                    }
                    for row in report.get("benchmarks", [])
                ],
                "key_gaps": top_project_gaps(report, limit=6),
                "plan_text": project.get("text"),
            }
        )
    with tempfile.NamedTemporaryFile("w", delete=False, suffix=".json", encoding="utf-8") as handle:
        json.dump(exportable, handle, indent=2)
        return handle.name


def qwen_project_prompt(project: dict[str, Any], extra_instruction: str) -> str:
    report = project.get("report", {})
    benchmark_lines = [
        f"- {row[0]}: {row[1]} ({row[2]}). Required action: {row[5]}"
        for row in report.get("benchmarks", [])
    ]
    checklist_lines = [
        f"- {item['criterion']}: {item['status']}. Recommendation: {item['recommendation']}"
        for item in report.get("checklist", [])
        if item.get("status") != "Reported"
    ]
    return f"""
You are OpenStudy's research benchmark review agent.

Review only the supplied project text and benchmark results. Do not invent approvals, methods, outcomes, or safeguards. Do not certify that the study is ethical or unbiased. Identify concrete protocol improvements a scientist can make before or during the project.

Project title: {project.get('title', 'Untitled project')}
Study type: {report.get('study_label', 'Unknown')}
Overall benchmark screen: {report.get('overall', 0):.1f}/100 - {report.get('level', 'Unknown')}

Benchmarks:
{chr(10).join(benchmark_lines) or '- No benchmark rows available.'}

Open checklist gaps:
{chr(10).join(checklist_lines[:16]) or '- No open checklist gaps flagged.'}

Project text:
{truncate(project.get('text', ''), MAX_QWEN_CONTEXT_CHARS)}

Extra reviewer instruction:
{normalize_space(extra_instruction) or 'None'}

Return Markdown with these sections:
1. Qwen benchmark verdict
2. Highest-priority fixes before continuing
3. Bias and ethics risks to monitor
4. Evidence missing from the project file
5. Suggested next protocol edits
"""


def qwen_agent_review(project: dict[str, Any], hf_token: str, extra_instruction: str) -> str:
    token = normalize_space(hf_token) or os.getenv("HF_TOKEN") or os.getenv("HUGGINGFACEHUB_API_TOKEN")
    prompt = qwen_project_prompt(project, extra_instruction)
    messages = [
        {
            "role": "system",
            "content": (
                f"You are a research-methods benchmark reviewer. You must use only {QWEN_MODEL_ID} "
                "as the configured model and must not mention or rely on any other model."
            ),
        },
        {"role": "user", "content": prompt},
    ]
    try:
        client = InferenceClient(model=QWEN_MODEL_ID, token=token or None, timeout=90)
        response = client.chat_completion(
            messages=messages,
            model=QWEN_MODEL_ID,
            max_tokens=900,
            temperature=0.2,
        )
        content = response.choices[0].message.content
        if not content:
            raise RuntimeError("The model returned an empty response.")
        return f"### Qwen agent review\n\n**Model used:** `{QWEN_MODEL_ID}`\n\n{content}"
    except Exception as exc:
        return (
            "### Qwen agent review unavailable\n\n"
            f"**Required model:** `{QWEN_MODEL_ID}`\n\n"
            "OpenStudy did not run another model. Configure a Hugging Face token as a Space secret named `HF_TOKEN`, "
            "or paste a temporary token in the project tab, then retry.\n\n"
            f"Technical detail: `{truncate(str(exc), 500)}`"
        )


def qwen_agent_review_ui(
    choice: str,
    projects: list[dict[str, Any]] | None,
    hf_token: str,
    extra_instruction: str,
) -> str:
    projects = projects or []
    if not projects:
        return "### Qwen agent review\n\nAdd or upload a research project first."
    if not choice:
        index = len(projects) - 1
    else:
        match = re.match(r"^(\d+)\.", choice)
        index = int(match.group(1)) - 1 if match else len(projects) - 1
    if index < 0 or index >= len(projects):
        return "### Qwen agent review\n\nSelected project was not found in this session."
    return qwen_agent_review(projects[index], hf_token, extra_instruction)


def extract_text_from_file(path: str) -> str:
    suffix = Path(path).suffix.lower()
    try:
        if suffix == ".pdf":
            reader = PdfReader(path)
            pages = []
            for page in reader.pages[:80]:
                pages.append(page.extract_text() or "")
            return normalize_space("\n".join(pages))
        if suffix == ".docx":
            document = Document(path)
            blocks = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    blocks.append(" | ".join(cell.text for cell in row.cells))
            return normalize_space("\n".join(blocks))
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        return f"Could not extract file text: {exc}"


def report_outputs(report: dict[str, Any]):
    return (
        summary_markdown(report),
        score_plot(report["scores"]),
        conduct_markdown(report),
        report.get("benchmarks", []),
        report["biases"],
        checklist_rows(report),
        write_report_file(report),
    )


def empty_outputs(message: str):
    fig, ax = plt.subplots(figsize=(7, 2.5))
    style_plot(fig, ax)
    ax.text(0.5, 0.5, message, ha="center", va="center", wrap=True, color=INK_SOFT)
    ax.axis("off")
    return (
        f"### {message}",
        fig,
        "### How the study appears to have been conducted\nNo study has been evaluated yet.",
        [],
        [],
        [],
        None,
    )


def build_app() -> gr.Blocks:
    blocks_kwargs: dict[str, Any] = {"title": f"{APP_NAME}: Study Bias and Ethics Screen"}
    launch_params = inspect.signature(gr.Blocks.launch).parameters
    if "theme" not in launch_params:
        blocks_kwargs["theme"] = THEME
    if "css" not in launch_params:
        blocks_kwargs["css"] = APP_CSS

    with gr.Blocks(**blocks_kwargs) as demo:
        gr.HTML(
            f"""
<header class="os-hero">
  <div class="os-wordmark">
    <span class="os-logo" aria-hidden="true">
      <svg viewBox="0 0 24 24" width="17" height="17" fill="none" stroke="currentColor"
           stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round">
        <path d="M12 2l8 4v6c0 5-3.5 8.5-8 10-4.5-1.5-8-5-8-10V6z"/>
        <path d="M9 12l2 2 4-4"/>
      </svg>
    </span>
    <span>{APP_NAME}</span>
  </div>
  <h1>Screen any study for the safeguards it actually reports.</h1>
  <p class="os-lede">Search public study metadata, upload a manuscript, or review your own protocol.
  {APP_NAME} checks for reported bias controls, study-conduct details, ethics safeguards, funding
  disclosures, and transparency signals &mdash; and tells you what is missing.</p>
  <div class="os-badges">
    <span class="os-badge os-badge-brand">Rule-based &amp; transparent</span>
    <span class="os-badge">Nothing stored after your session</span>
    <span class="os-badge">Optional AI review: {QWEN_MODEL_ID}</span>
    <span class="os-badge">Open source &middot; MIT</span>
  </div>
  <div class="os-note"><strong>Decision support, not a verdict.</strong> {APP_NAME} flags what is
  reported or missing in the supplied text; it does not certify that a study is ethical, unbiased,
  valid, or invalid.</div>
</header>
"""
        )

        with gr.Tab("Search studies"):
            with gr.Row():
                query = gr.Textbox(
                    label="Study title, DOI, URL, or keywords",
                    placeholder="Example: 10.1056/NEJMoa2034577 or paste a paper title",
                    scale=5,
                )
                rows = gr.Slider(1, 10, value=5, step=1, label="Results")
            search_button = gr.Button("Search public study metadata", variant="primary")
            search_status = gr.Markdown()
            studies_state = gr.State([])
            results_table = gr.Dataframe(
                headers=["#", "Title", "Year", "Venue", "DOI", "Source", "Available abstract preview"],
                datatype=["number", "str", "str", "str", "str", "str", "str"],
                interactive=False,
                wrap=True,
                elem_classes=["compact-table"],
            )
            selected = gr.Dropdown(label="Select a result to evaluate", choices=[], interactive=True)
            analyze_button = gr.Button("Evaluate selected study", variant="secondary")

            with gr.Row():
                summary = gr.Markdown()
                plot = gr.Plot(label="Category scores", show_label=False)
            conduct = gr.Markdown()
            benchmarks_table = gr.Dataframe(
                headers=["Benchmark", "Status", "Score", "Description", "Evidence", "Required action"],
                interactive=False,
                wrap=True,
            )
            bias_table = gr.Dataframe(
                headers=["Signal", "Status", "Evidence from available text", "Suggested next check"],
                interactive=False,
                wrap=True,
            )
            checklist = gr.Dataframe(
                headers=["Category", "Criterion", "Status", "Evidence", "Recommendation"],
                interactive=False,
                wrap=True,
            )
            report_file = gr.File(label="Download Markdown report")

            search_button.click(
                search_studies_ui,
                inputs=[query, rows],
                outputs=[results_table, selected, studies_state, search_status],
            )
            analyze_button.click(
                analyze_selected_ui,
                inputs=[selected, studies_state],
                outputs=[summary, plot, conduct, benchmarks_table, bias_table, checklist, report_file],
            )

        with gr.Tab("Upload manuscript"):
            with gr.Row():
                upload = gr.File(
                    label="Upload study PDF, DOCX, TXT, or Markdown",
                    file_types=[".pdf", ".docx", ".txt", ".md"],
                    type="filepath",
                )
                with gr.Column():
                    upload_title = gr.Textbox(label="Study title (optional)")
                    type_hint = gr.Dropdown(
                        label="Study type hint",
                        choices=STUDY_TYPE_CHOICES,
                        value="Auto-detect",
                    )
            pasted = gr.Textbox(
                label="Paste abstract, methods, or manuscript text (optional)",
                lines=10,
                placeholder="Paste text here if you do not have a PDF or want to supplement extraction.",
            )
            upload_button = gr.Button("Evaluate uploaded study", variant="primary")
            with gr.Row():
                upload_summary = gr.Markdown()
                upload_plot = gr.Plot(label="Category scores", show_label=False)
            upload_conduct = gr.Markdown()
            upload_benchmarks = gr.Dataframe(
                headers=["Benchmark", "Status", "Score", "Description", "Evidence", "Required action"],
                interactive=False,
                wrap=True,
            )
            upload_bias = gr.Dataframe(
                headers=["Signal", "Status", "Evidence from available text", "Suggested next check"],
                interactive=False,
                wrap=True,
            )
            upload_checklist = gr.Dataframe(
                headers=["Category", "Criterion", "Status", "Evidence", "Recommendation"],
                interactive=False,
                wrap=True,
            )
            upload_report_file = gr.File(label="Download Markdown report")

            upload_button.click(
                analyze_upload_ui,
                inputs=[upload, pasted, upload_title, type_hint],
                outputs=[upload_summary, upload_plot, upload_conduct, upload_benchmarks, upload_bias, upload_checklist, upload_report_file],
            )

        with gr.Tab("My research projects"):
            gr.Markdown(
                """
Add a study idea, protocol, or active project to review whether the planned conduct follows the benchmarks readers and reviewers will expect.

Project data is kept in this browser session only. Download the report or portfolio export if you want to keep it.
"""
            )
            projects_state = gr.State([])
            project_file = gr.File(
                label="Upload project, protocol, benchmark plan, or preregistration",
                file_types=[".pdf", ".docx", ".txt", ".md", ".json", ".csv", ".tsv"],
                type="filepath",
            )
            with gr.Row():
                project_title = gr.Textbox(label="Project title", placeholder="Example: Community sleep intervention pilot")
                project_team = gr.Textbox(label="Team or owner", placeholder="Lab, PI, student group, or organization")
            with gr.Row():
                project_status = gr.Dropdown(
                    label="Project status",
                    choices=["Idea", "Protocol drafting", "Ethics review", "Recruiting", "Data collection", "Analysis", "Manuscript drafting", "Completed"],
                    value="Protocol drafting",
                )
                project_type = gr.Dropdown(label="Study type", choices=STUDY_TYPE_CHOICES, value="Auto-detect")

            research_question = gr.Textbox(label="Research question or hypothesis", lines=2)
            with gr.Accordion("Design and participants", open=True):
                study_design = gr.Textbox(label="Study design", lines=2, placeholder="Design, setting, timeline, and whether it is prospective, retrospective, randomized, observational, qualitative, etc.")
                participants = gr.Textbox(label="Population or sample", lines=2, placeholder="Who or what will be studied, expected sample, and important demographics.")
                recruitment = gr.Textbox(label="Recruitment source and study period", lines=2, placeholder="Where participants/data come from, how they are approached, and when data will be collected.")
                eligibility = gr.Textbox(label="Inclusion and exclusion criteria", lines=2)
            with gr.Accordion("Measures and bias controls", open=True):
                intervention_or_exposure = gr.Textbox(label="Intervention, exposure, or focus", lines=2)
                comparator = gr.Textbox(label="Comparator or control", lines=2, placeholder="Placebo, usual care, matched controls, pre/post comparison, or why none is appropriate.")
                outcomes = gr.Textbox(label="Primary and secondary outcomes", lines=2, placeholder="Name primary outcomes before data collection when possible.")
                sample_size = gr.Textbox(label="Sample size or power rationale", lines=2)
                bias_controls = gr.Textbox(label="Bias controls", lines=3, placeholder="Randomization, blinding, allocation concealment, validated measures, training, calibration, reflexivity, or confounder controls.")
            with gr.Accordion("Ethics, analysis, and transparency", open=True):
                ethics_plan = gr.Textbox(label="Ethics approval and oversight", lines=2, placeholder="IRB/ethics committee, exemption rationale, animal-care approval, or oversight plan.")
                consent_privacy = gr.Textbox(label="Consent, privacy, and safety safeguards", lines=2)
                analysis_plan = gr.Textbox(label="Statistical or qualitative analysis plan", lines=3)
                missing_data_plan = gr.Textbox(label="Missing data, attrition, or follow-up plan", lines=2)
                transparency_plan = gr.Textbox(label="Data, code, materials, and preregistration plan", lines=2)
                funding_conflicts = gr.Textbox(label="Funding, sponsor role, and conflicts of interest", lines=2)
                project_notes = gr.Textbox(label="Additional notes", lines=2)

            add_project_button = gr.Button("Add project and review plan", variant="primary")
            project_status_text = gr.Markdown()
            project_table = gr.Dataframe(
                headers=["#", "Project", "Detected/review type", "Status", "Score", "Risk level", "Key gaps"],
                datatype=["number", "str", "str", "str", "str", "str", "str"],
                interactive=False,
                wrap=True,
            )
            with gr.Row():
                project_select = gr.Dropdown(label="Select a saved project from this session", choices=[], interactive=True)
                review_project_button = gr.Button("Review selected project", variant="secondary")
                export_projects_button = gr.Button("Download project portfolio JSON")
            projects_file = gr.File(label="Project portfolio export")

            with gr.Accordion(f"Qwen agent benchmark review ({QWEN_MODEL_ID} only)", open=False):
                gr.Markdown(
                    f"""
The AI reviewer is constrained to `{QWEN_MODEL_ID}`. OpenStudy will not fall back to another model.

For Hugging Face Spaces, set a secret named `HF_TOKEN`, or paste a temporary Hugging Face token below for this session.
"""
                )
                qwen_token = gr.Textbox(label="Hugging Face token (optional)", type="password")
                qwen_instruction = gr.Textbox(
                    label="Extra Qwen review focus (optional)",
                    lines=2,
                    placeholder="Example: Focus on recruitment fairness and IRB readiness.",
                )
                qwen_button = gr.Button("Run Qwen benchmark agent", variant="secondary")
                qwen_output = gr.Markdown()

            with gr.Row():
                project_summary = gr.Markdown()
                project_plot = gr.Plot(label="Category scores", show_label=False)
            project_conduct = gr.Markdown()
            project_benchmarks = gr.Dataframe(
                headers=["Benchmark", "Status", "Score", "Description", "Evidence", "Required action"],
                interactive=False,
                wrap=True,
            )
            project_bias = gr.Dataframe(
                headers=["Signal", "Status", "Evidence from project plan", "Suggested next check"],
                interactive=False,
                wrap=True,
            )
            project_checklist = gr.Dataframe(
                headers=["Category", "Criterion", "Status", "Evidence", "Recommendation"],
                interactive=False,
                wrap=True,
            )
            project_report_file = gr.File(label="Download project review report")

            project_inputs = [
                project_file,
                project_title,
                project_team,
                project_status,
                project_type,
                research_question,
                study_design,
                participants,
                recruitment,
                eligibility,
                intervention_or_exposure,
                comparator,
                outcomes,
                sample_size,
                bias_controls,
                ethics_plan,
                consent_privacy,
                analysis_plan,
                missing_data_plan,
                transparency_plan,
                funding_conflicts,
                project_notes,
                projects_state,
            ]
            add_project_button.click(
                add_project_ui,
                inputs=project_inputs,
                outputs=[
                    project_table,
                    project_select,
                    projects_state,
                    project_status_text,
                    project_summary,
                    project_plot,
                    project_conduct,
                    project_benchmarks,
                    project_bias,
                    project_checklist,
                    project_report_file,
                ],
            )
            review_project_button.click(
                review_project_ui,
                inputs=[project_select, projects_state],
                outputs=[project_summary, project_plot, project_conduct, project_benchmarks, project_bias, project_checklist, project_report_file],
            )
            export_projects_button.click(download_projects_ui, inputs=[projects_state], outputs=[projects_file])
            qwen_button.click(
                qwen_agent_review_ui,
                inputs=[project_select, projects_state, qwen_token, qwen_instruction],
                outputs=[qwen_output],
            )

        with gr.Tab("Method"):
            gr.Markdown(
                """
## How the screen works

OpenStudy uses rule-based, transparent benchmark checks inspired by common reporting safeguards across trials, observational studies, reviews, qualitative research, animal studies, and model studies.

The score is a **reported-or-planned safeguard score**, not a truth score. A low score can mean the study is poorly reported, the public abstract is too short, the uploaded text did not include the methods, ethics, funding, and limitations sections, or a project plan still needs those safeguards added.

The optional AI review uses **`Qwen/Qwen3.6-27B` only** through Hugging Face inference. If that model is not available with the configured token/provider, OpenStudy does not switch to another model.

Recommended human follow-up:

- Read the full methods, supplement, protocol, registry entry, and conflict-of-interest statements.
- Compare outcomes in the manuscript with the protocol or registration.
- Ask domain experts to judge whether the design and analysis fit the research question.
- Treat ethics and participant-protection findings as prompts for review, not legal or IRB determinations.
- For unpublished projects, use the project report as a protocol-improvement checklist before recruitment or data collection.
- Upload protocol, benchmark, or preregistration files in the project tab to check whether planned safeguards are present.
"""
            )

        gr.HTML(
            f"""
<footer class="os-footer">
  <span>{APP_NAME} &mdash; open-source study screening &middot; MIT license</span>
  <span>A reported-safeguard screen, not a verdict on validity or ethics.</span>
</footer>
"""
        )

    return demo


demo = build_app()


def launch_app() -> None:
    launch_kwargs: dict[str, Any] = {}
    launch_params = inspect.signature(demo.launch).parameters
    if "theme" in launch_params:
        launch_kwargs["theme"] = THEME
    if "css" in launch_params:
        launch_kwargs["css"] = APP_CSS
    demo.launch(**launch_kwargs)


if __name__ == "__main__":
    launch_app()
